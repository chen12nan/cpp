from docx import Document
from os import listdir
from docx.shared import Inches
document = Document()
pictures = [x for x in listdir() if x.endswith('.jpg')]
for pic in pictures:
    print(pic)
    try:
        document.add_picture(pic,width=Inches(1),height=Inches(2))
    except:
        print('can not konw the style',pic)
# document.add_picture('a.jpg')

document.save('world.docx')

# document.add_paragraph('Hello,Word!')
# document.save('demo.docx')